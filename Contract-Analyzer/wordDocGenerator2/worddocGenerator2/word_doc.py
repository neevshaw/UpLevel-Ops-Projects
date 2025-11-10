from datetime import datetime, timezone
import io
import base64
import copy

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'lib')) # Prepare to import from the lib directory

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.text.run import Run

def current_document_datetime():
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

def wrap_runs_with_tracked_deletion(runs, author, deletion_id):
	if not runs:
		return None

	paragraph = runs[0]._element.getparent()
	first_run_element = runs[0]._element
	insertion_point = paragraph.index(first_run_element)

	# Create the <w:del> element
	del_element = OxmlElement("w:del")
	del_element.set(qn("w:id"), str(deletion_id))
	del_element.set(qn("w:author"), author)
	del_element.set(qn("w:date"), current_document_datetime())

	# Clone each run and replace <w:t> with <w:delText>
	for run in runs:
		run_clone = copy.deepcopy(run._element)
		for t in run_clone.findall(".//w:t", namespaces=run_clone.nsmap):
			t.tag = qn("w:delText")
		del_element.append(run_clone)

	# Remove the original runs
	for run in runs:
		paragraph.remove(run._element)

	# Insert the deletion at the original position
	paragraph.insert(insertion_point, del_element)

	return del_element

def create_tracked_deletion(prev_element, runs, author, change_id):
    change_date = current_document_datetime()
    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), change_date)
    del_elem.set(qn('w:id'), str(change_id))

    for run in runs:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        styles = run.get('styles', [])
        if 'bold' in styles:
            rPr.append(OxmlElement('w:b'))
        if 'italic' in styles:
            rPr.append(OxmlElement('w:i'))
        if 'underline' in styles:
            rPr.append(OxmlElement('w:u', {'w:val': 'single'}))
        if len(rPr):
            r.append(rPr)

        del_text = OxmlElement('w:delText')
        del_text.text = run['text']
        r.append(del_text)
        del_elem.append(r)

    prev_element.addnext(del_elem)
    return del_elem

def create_tracked_insertion(prev_element, runs, author, change_id):
    change_date = current_document_datetime()
    ins_elem = OxmlElement('w:ins')
    ins_elem.set(qn('w:author'), author)
    ins_elem.set(qn('w:date'), change_date)
    ins_elem.set(qn('w:id'), str(change_id))

    for run in runs:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        styles = run.get('styles', [])
        if 'bold' in styles:
            rPr.append(OxmlElement('w:b'))
        if 'italic' in styles:
            rPr.append(OxmlElement('w:i'))
        if 'underline' in styles:
            rPr.append(OxmlElement('w:u'))

        # Handle size-XX styles (e.g. size-12)
        for style in styles:
            if style.startswith('size-'):
                size = style.split('-')[1]
                sz_elt = rPr._add_sz()
                sz_elt.val = int(float(size) * 914400 / 72)
                break

        if len(rPr):
            r.append(rPr)

        t = OxmlElement('w:t')
        t.text = run['text']
        r.append(t)
        ins_elem.append(r)

    prev_element.addnext(ins_elem)
    return ins_elem

# Register comment in comments.xml
def add_comment_to_change(change_elem, comment_id):
    start = OxmlElement('w:commentRangeStart')
    start.set(qn('w:id'), str(comment_id))
    change_elem.addprevious(start)

    end = OxmlElement('w:commentRangeEnd')
    end.set(qn('w:id'), str(comment_id))
    change_elem.addnext(end)

    ref_run = OxmlElement('w:r')
    ref = OxmlElement('w:commentReference')
    ref.set(qn('w:id'), str(comment_id))
    ref_run.append(ref)
    end.addnext(ref_run)

# Insert comment-related elements into document.xml
def add_comment(document, comment_id, author, text):
    comments_part = document.part._comments_part
    if comments_part is None:
        comments_part = document.part._add_comments_part()
        comments = parse_xml(r'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        comments_part._element = comments

    comments_el = comments_part._element

    comment = OxmlElement('w:comment')
    comment.set(qn('w:author'), author)
    comment.set(qn('w:date'), current_document_datetime())
    comment.set(qn('w:id'), str(comment_id))

    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    comment.append(p)

    comments_el.append(comment)

# Make a duplicate of a run immediately after it
def duplicate_run(run):
    r_elt = run._element

    new_elt = copy.deepcopy(r_elt)
    new_run = Run(new_elt, run._parent)

    r_elt.addnext(new_elt)

    return new_run

# Split a run before the provided index
def split_run(run, index):
    new_run = duplicate_run(run)

    # Remove text from the old and new run
    run.text = run.text[:index]
    new_run.text = new_run.text[index:]

    return (run, new_run)

def doc_from_json_doc_data(doc_data):
    document = Document()
    change_id_counter = 1

    for paragraph in doc_data['paragraphs']:
        p = document.add_paragraph()
        for section in paragraph:
            section_type = section['type']
            if section_type == 'edit':
                author = section.get('author', 'Unknown Author')
                comment_text = section.get('comment', {}).get('text')

                comment_id = change_id_counter
                if comment_text:
                    add_comment(document, comment_id, author, comment_text) # Register comment in comments.xml

                del_elem = None
                ins_elem = None

                if 'old_runs' in section and section['old_runs']:
                    prev_element = p.add_run()._r
                    del_elem = create_tracked_deletion(prev_element, section['old_runs'], author, change_id_counter)
                    change_id_counter += 1

                # Only add the comment to the new section
                if 'runs' in section and section['runs']:
                    prev_element = p.add_run()._r
                    ins_elem = create_tracked_insertion(prev_element, section['runs'], author, change_id_counter)
                    change_id_counter += 1
                
                # Add comment to document.xml
                if comment_text:
                    if ins_elem is None: # If this is a deletion with no insertion
                        add_comment_to_change(del_elem, comment_id)
                    else:
                        add_comment_to_change(ins_elem, comment_id)
            else:
                for run in section['runs']:
                    styles = run.get('styles', [])
                    r = p.add_run(run['text'])
                    if 'bold' in styles:
                        r.bold = True
                    if 'italic' in styles:
                        r.italic = True
                    if 'underline' in styles:
                        r.underline = True
                    if 'strikethrough' in styles:
                        r.font.strike = True

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    encoded = base64.b64encode(buffer.read()).decode('utf-8')
    return encoded

def apply_list_formatting(new_paragraph, doc, list_level=0):
    from copy import deepcopy
    nsmap = new_paragraph._element.nsmap
    
    # Try to find an existing paragraph with the same level
    for p in doc.paragraphs:
        ppr = p._element.pPr
        if ppr is not None:
            numPr = ppr.find('.//w:numPr', nsmap)
            if numPr is not None:
                ilvl = numPr.find('.//w:ilvl', nsmap)
                if ilvl is not None and int(ilvl.get('{%s}val' % nsmap['w'])) == list_level:
                    new_paragraph._element.insert(0, deepcopy(ppr))
                    return
    
    # If not found, create a new numPr with desired level
    from lxml import etree
    numPr = etree.Element('{%s}numPr' % nsmap['w'])
    ilvl = etree.Element('{%s}ilvl' % nsmap['w'])
    ilvl.set('{%s}val' % nsmap['w'], str(list_level))
    numId = etree.Element('{%s}numId' % nsmap['w'])
    numId.set('{%s}val' % nsmap['w'], "1")  # assumes list definition 1
    
    numPr.append(ilvl)
    numPr.append(numId)

    if new_paragraph._element.pPr is None:
        new_paragraph._element.insert(0, etree.Element('{%s}pPr' % nsmap['w']))

    new_paragraph._element.pPr.append(numPr)

def json_data_from_doc(doc, cursor, size):
    cursor = int(cursor)
    size = int(size)
    paragraphs = []
    for paragraph in doc.paragraphs[cursor:min(cursor+size, len(doc.paragraphs))]:
        runs = []

        namespaces = {'w': "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        # Find all run elements
        for r_elt in paragraph._element.findall(".//w:r", namespaces=namespaces): # Simply doing paragraph.runs() will not account for insertions and deletions
            # Determine whether this run is inside an insertion, a deletion, or neither
            is_insertion = False
            is_deletion = False
            parent_elt = r_elt.getparent()
            if parent_elt is not None:
                if parent_elt.tag == qn('w:ins'):
                    is_insertion = True
                elif parent_elt.tag == qn('w:del'):
                    is_deletion = True

            run = Run(r_elt, paragraph._element) # Create a run object so styles are easily accessible
            text = run.text
            # Correct run text if the run contains a delText (Run initialization doesn't work properly in that case)
            del_text = r_elt.findall(".//w:delText", namespaces=namespaces)
            if len(del_text) > 0:
                text = ''
                for del_text_item in del_text:
                    text += del_text_item.text

            styles = []
            if run.bold:
                styles.append('bold')
            if run.italic:
                styles.append('italic')
            if run.underline:
                styles.append('underline')
            if run.font.strike:
                styles.append('strikethrough')
            
            rPr = r_elt.rPr
            if rPr is not None:
                sz_element = rPr.sz
                if sz_element is not None:
                    size_emu = int(sz_element.val) # Size in EMU
                    font_size_pt = size_emu / 914400 * 72
                    styles.append(f'size-{round(font_size_pt, 1)}')

            new_run_data = {
                'text': text,
                'styles': styles
            }
            # Mark this run as an insertion or deletion if appropriate
            if is_insertion:
                new_run_data['is_insertion'] = True
            elif is_deletion:
                new_run_data['is_deletion'] = True
            runs.append(new_run_data)

        is_list_item = paragraph._element.find('.//w:numPr', namespaces=namespaces) is not None

        paragraphs.append([{
            'type': 'normal',
            'is_list_item': is_list_item,
            'runs': runs
        }])
    
    return {'doc_data': {'paragraphs': paragraphs}, 'next-cursor': (cursor+ size if cursor + size < len(doc.paragraphs) else "STOP")}