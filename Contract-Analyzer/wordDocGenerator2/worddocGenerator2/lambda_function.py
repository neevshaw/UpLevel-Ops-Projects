import json
import io
import base64
import boto3
import uuid
import cgi
import re
import os

from docx import Document
from docx.text.run import Run

# Import helper functions from word_doc.py
from word_doc import (
    json_data_from_doc,
    split_run,
    wrap_runs_with_tracked_deletion,
    create_tracked_insertion,
    doc_from_json_doc_data
)

TMP_DATA_UPLOAD_BUCKET_NAME = 'tmp-word-doc-json-upload'
TMP_DOC_UPLOAD_BUCKET_NAME = 'tmp-word-doc-upload'

FAVICON_ELT = '<link rel="icon" href="https://uplevelops.com/wp-content/uploads/2024/09/cropped-ULO22_Symbol-White-R-on-blue-32x32.png" sizes="32x32">'
LOGO_ELT = '<img id="logo" src="data:image/jpeg;base64,/9j/4AAQSkZJRgABAQAAAQABAAD/2wBDAAMCAgMCAgMDAwMEAwMEBQgFBQQEBQoHBwYIDAoMDAsKCwsNDhIQDQ4RDgsLEBYQERMUFRUVDA8XGBYUGBIUFRT/2wBDAQMEBAUEBQkFBQkUDQsNFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBT/wgARCADIAMgDASIAAhEBAxEB/8QAHQABAAIDAAMBAAAAAAAAAAAAAAcIBQYJAQMEAv/EABsBAQEBAAMBAQAAAAAAAAAAAAAFBAECBgMH/9oADAMBAAIQAxAAAAG1JCPf5zchF9Ok3IR3Hrzvo+X2AAPxhun1zjTXw2bk00bk00bk03288ba8edM8Bzj6Oc46UwKUxP8AAE/59NnRDugAAatFk+apL9JE4g+2Ht5Jce294bOtW2mn54Bzj6Oc46UwKUxP8AT/AJ9NnRDugAAAQhi9s1jyn6eltsdjyStk3x7U85qlgIgm3jn9gc4+jnOOlMClMT/AE/59NnRDugAAD5+OYs3j2Z/DaVu99bfQ+RsrYevFh8uoM+kBzj6Oc46UwKUxP8AT/n02dEO6AAA8eQr776v0ZvgU5dlbD14sPDuhn0gOcfRznHSmBSmJ/gCf8+mzoh3QAAGKyuH568/PB6PzYFlbD14sPDuhn0gOcfRznHSmBSmJ/gCf8+mzoh3QAAGHzGH568+x6PzYFlbD14sPDuhn0gOcfRznHSmBSmJ/gCf8+mzoh3QAAGHzGH568+x6PzYFlbD14sPDuhn0gOcfRznHSmBSmJ/gCf8APps6Id0AABh8xq/brQsei82BZWw9eLDw7oZ9IDnH0c590ZuHFOWn2Apj+Gi3IhXgAAFTJagLfP0b4rlU13YA+/wsrYevFh4d0M+kBUq2up/fPQ99Xy3ITYdeOejX6heaPP8AoA6fQBoeArJsxfq6eK3br2U7uJrXTvQh9Xy3IVlbD14sPEuBn0gARTUXohq2zFQxNES1JXqsfWdx2vl76De3Jrt9B3yTlxzAluNk/WTWGfSBXCuXRKkdWTL9h4GnnFtDPpAAAeoNZxp3+f2bMcdvrHXsAAAAAAB//8QAKhAAAQIEBgEEAwEBAAAAAAAABQMEAgYHMAABECAzNBYRFBc1EhMVJUD/2gAIAQEAAQUC/wCr5ZFY+WRWPlkVj5ZFYlucWkzq7os88sOSntcvKW2PKW2PKW2PKW2PKW2PKW2PKW2EpkQWUsUi79gkBSeZLIxt1NiSUSyggRCOTsUi79k+Nydt9UkollBAiEemVOMAkAuZxZlbfSLv2ibf2r/CacSyggRCPTw6eLJDolnB4SNdZvh26kXftTLl6Ek04loxAiEfBgWebGHZQE4HulhjuY3MEEKUG6kXftG/V2XECIR8GJ+nv8sUh4bNIu/ZXV/SkMGZNNJ+nv8ADSkPDZpF37PppPs9+31pDw2aRd+1Ps9eyy1pDw2aRd+yVjzTF55+uetIeGzSLv2TH1GykPDZpF37Jj6jZSHhs0i79kx9RspDw2aRd+yY+o2Uh4bNIu/ZMfUbKQ8NmkXfsmPqNlIeGzSLv2ZnfJDwGykPDZpFn/oWalTL/TIOWS7PZSHh2mGubEtrSxz+mZLE/TbkCZSTLMUxlZ+lrI0G1pDw7aoiM2ZvWXyX8g1DFlFDumya0JZZtm7+bTQIKgAG4qHLf8UtpSHh2zOBgmIQ6aqsnGtOJohJD9s11Aag8kUSM1lpWlZvLLLSYgqcwCXLZRm4xSHh3TrJMExJu2izFxo3cqtFwFVEo4G81B3UKsyCkciVTA7LI9UMmZyluTH8xxgZdZy612VRlvSkPDvOS0wmFEzS8iyzdMXDGPVNKNaMXT8yTzB0zHjc4YcoIdrpqm9bGJSIDCVKmLhkjZUSgWhXlQO4x4IBwlJ4VDDdog0hv//EAC4RAAECBAMGBgIDAAAAAAAAAAECBAADEDMRIEIFEhMiMTIUFSEjUVOhsTBhcf/aAAgBAwEBPwFS0o7jHHl/MCahRwByS5SppwRHlzv6zHlzv6zHlzv6zB2e6HqUVeaaN7oy7N2kqWoSZp5T+KE4epjaW0vEe1K7f3V5po3ujM1mb7dC1fEbS2lx/aldv7gEHpV5po3ujKBicBDt9jLDaSeUfmHDjQiG9oVeaaN7ozOHGHIije0KvNNG90ZT0q3tCrzTRvdGU9Kt7Qq800b3RlPSre0KvNNG90ZVdDVvaFXmmkg4TBlcLKjwkROlcJWFG9oVnS+IjCgOBxhCwtO8Kzp+7yo6xIk7nMrrE6XxE4Ub2hknN+JzJ6wqWpHcIlzVS+2PGK+I9+d/QiVITK/2rmVgreEN7YzbifiAkDoP4f/EACsRAAAFAgUDBAIDAAAAAAAAAAECAwQQADMREiAxQhMUIiEyUVIjMEFhcf/aAAgBAgEBPwEpDH9oV0FPihSOUMRDQY4F3ruUftXco/au5R+1dyl9pZ8ocWx0umoGDOTeWrXp+Z95Z8ocWx1KlyqCUKaten5n3oQEN5Z8ocWx0j6UihgYVT7jTdDmanF0ZZ8ocWx1N0OZocXRlnyhxbHSG8uLoyz5Q4tjpDeXF0ZZ8ocWx0hvLi6Ms+UOLY6S7hLi6Ms/5hf1THSgQCh1TUkr1S4w4ujKJ+mfGBDEMKOQSGyjKSGbyPtS62fxLtSKnTNjDi6OhFxk8TbUU5T+0aOmVT3V2hfmvwI/2NKrGU/yWyuIZRpxcHVnN80JhHf9P//EAEUQAAECAQUKCQoFBAMAAAAAAAIBAwQABRESMBMgITFBUXJzscEQFCIyM2FxkdEjJDRCYoGSobLSQENSguEGFTVTY4Oi/9oACAEBAAY/AvxXosZ8I/dL0WM+Eful6LGfCP3S9FjPhH7pPNwzTzatIhLdUTct/gSmVJwr9XOiIqbZdG73J4y6N3uTxl0bvcnjLo3e5PGXRu9yeMujd7k8ZdG73J4yQAZeIlxJQnjYzhqx22Km0iNPdWJZE24NUkyXqACViXEiSrFyn1xrmsZw1Y7bJXQTyrfzS8QASsS4kSVYuU+uNc0hKNigYrc1F5xdiY1krMJFibyYbkSKB9xYbCcNWO2zebTEi4OBABKxLiRJVi5T641zcE4/1AL1yj3Jw4qD1zRxWWUOrVFFRetZT5EnErFf2rysFOJMXF1HBGsQ0d2TLKFiFSqrzQuUZqUpv5w1Y7bPtBJIAJWJcSJKsXKfXGubgjGYZboMKqCTiYlVacXdKJucEc5zTExCRZQ7JoLrLyYaUp5yLRilFgxN0RM8HHKPHXok0ruIOQQRVoVcVMhAUqiKUIl/OGrHbZkDaVlwCiJKsXKfXGubgcmybnMGJ58fpSU56Te+ynDVjtslKisuRM6yJ1zlxJ4SLNwOTZNznKxPPjk9lOCc9JvfZThqx22jk2zc55XmvPD6vsp18M56Te+ynDVjtszm2bj84xOvD6nUnXeTnpN77KcNWO2yjCFapIyaoqZMF7Oek3vspw1Y7bKO1B/TeznpN77KcNWO2yjtQf03s56Te+ynDVjtso7UH9N7Oek3vspw1Y7bKO1B/TeznpN77KcNWO2yjtQf03s56Te+ynDVjtso7UH9N7Oek3vspw1Y7bKOddWgbkQp1quBL2c9JvfZTgn/ABDtskgGCphoZeUqesf8SauzZN3UEcCnKK5byc9JvffRkOv5bpD87wm1XpmCH34F3WPFoc/PnkwUfljnkl0ReJs8p0s/s++VdgPOoVKzaJlHKN5Oek3vvhjBTyUUP/pMe68g4v1W3EraOX5SRUWlFypf0rQ5Fmnkmd69UqtKvRT5Uka4kTP2SbhGEwDhI8pln4OMMjRCRXKSj1SypwznpN7752FLA5zmz/SUnGHgVt1tapCuS8Gbnz86h0oGn1w/i+KHhVGKjsVCc0O3wlQlaKi3VpUlyeCSqB5SIPpXs/8AHC9CHgUsIF+ksiycYeGo62VUhXIvBOek3vv+MQ9DceCY8jiZlkbEQ2TToYFEuEHmTVp0FpEhxpIWp1bUD/3tJSi9qSpbnKG7CcQV+cqTnKFT/uSS3EzjDzNDQnesibbXiUOvqNLhXtKSEA3CFyxBpg92eVxhQwrz3C5x9t6M7sD7D6J8i3d3BOek3vsKkW1Sac10cBj75KUEQxzWbmn3SqRLDjBZnBovEFsFMlyClMkXi3FW/wBcRyfljkLkWvH3kyElAJ7sskEUQRTEiXzjDw12nBqknVJ+HCFffbFeQ4DaqhJklOPGGHWKxBRdAUacdlVMUMcxJTLlzbDftbRNkv8AHB8ReMuTNkP+4a22VVhltkczY0fgP//EACgQAAEBBwMEAwEBAAAAAAAAAAERACAhMDFR8EFhgRBxkaGxwdFA4f/aAAgBAQABPyH+6FChQi2akEEEpBR8LF2ghiGRn+aSDhw4cOHDgi40BEYCoChNpGGuk3qCDyj7Y5AtCTpN5aahoEAfzCThrpUApLDmDhN5aahoGAfzBgvs0XkmyvAGCJ9SOXSBDiRhrpYRJFobGI+eh55aahoGAfzDod1usH5CKIIktCEgePBQgIIBEF9FA6V4KPb+GuljUuPlj2S01DQOA/iHRW+EteC5NTHwyOWAS4AoDFVkSINYFDwRGaXYZMINQAUD+Gulkbo7oKNA4D+IdNyU7W5fk8NnLSuGulIGQYDroBlVgk7hoOm4Id5Fvc8dM5aVw10ogSCREU6JNWEuZfY6d6dM5aVw10uKSUe5m9jp3oSpUxPXOWlcNdKNHLCQkDQsQxJUmJJczlpXDXSsjc7nLSuGulZG53OWlcNdKyNzuctK4a6Vkbnc5aVw10rI3O5y0rhrpWRudzlpXDXSoCeLUNAcl3OWe0c7kJKiWUVkKZ+ndW0lqu0BzOWejAQ8cCRwacAALgfhKSRiQ6kdTvt50YjLgetworIoKAOQAo3G7mcs9grUSRoMAPCvLiykgoi6HuLATDKCQIfLiF0SvwB7ZFuSNRWAfQZfdEQiKnlE6LC0FFDHI77dc5Z6aIVto0PbQ7FjNi1Mg4KAI1RBRN6OybvahoBT7xU7OUYjsSUbnQDQlBAkETsLBbqmOEv08OxLG8kNIK9M5Z/QvwFD5RY4C4WR0I6rL4UhGAacOaqo4XsGGoh+IILIcXdPhWDQvQhTdD0rBbwxKS8nhGI5V+qdXrdoBcaKNcvqjqRqKBOPudGctIpAgJ4I+hg1d8QKB5QPB4bf59P25SScYR4DKSE1IcLMePFRSyr4YeEEAQAPCmkjdSaPFW+YkoFUruwJKIdIUrKIT5UQGLkqjUkfgxIVbB5SM2ptgAEfX8H/2gAMAwEAAgADAAAAEP8Az/zzznPPPHzz32jzzzx7/wAj8899o8888sU60A899o8888/Ut+8899o88888h9+8899o8888/wDffvPPfaPPPPP/AH37zz32jzzzz/337zz32jzzzzz337zzz3jzzzzjX37zz37zfzzrbvb7zzw7dw4UfzxTXzzzzx/zzzzzzzzz/8QAKBEAAQIFAwMFAQEAAAAAAAAAARARADGxwfAgIXFBUdFhgZGh4TDx/9oACAEDAQE/EGzYePRwOBEnQxJz2cP8ExgDzGAPMYA8wFJIA48resmHjS/y2ATMum/anCAEkYCDkTsMziVVvWTDxqKdzAk+25g5k7DM4lWALk4W9ZMPGkwphg23mAT3AUf5iYTk2itqVvWTDxq3BN+pslbUresmHjTOWtqVvWTDxpnLW1K3rJh40zlralb1kw8aSAMey1tSoli5sjjPfSFTjPKwTqAZJW1KvATmIIILGCCCYgS6ih6s/r9hjMPP9h/CYlBDFjFbU6BlI+hg0zEE3L26QJDXh9ux+xuA37lOx7GfP7AEDPrXUTzH4ESMH8f/xAAmEQABAgQHAQADAQAAAAAAAAABEBEAMbHBICFRYXHh8EEwkaHR/9oACAECAQE/EHzMaN3B4yAwA3NhwY2kbSNpAKWALYunnzhJxMExr3VACSwgAAG0adrYunnzibZ9LQAADaNO4IsDFbF08+cJACTDas4kDR7xJDwLxQ0C2Lp584nWDl8F0oaBbF08+cMpaGgWxdPPnDKWhoFsXTz5wyloaBbF08+cIEiGq0NApBx4ugMxphOpIl6kC+AiaUNArYUpGAQQ4gBikYIvgp/iD+9Q5kDDWUjOAXzEUNBgIMyiBzuQPYP9gvO9obMudRkpy0KJ3zEuOoIE42oMQHI/2YmYn8P/xAAoEAEAAQMDAwQCAwEAAAAAAAABEQAhMUFRYSBxgRAwQJGx0aHB8HD/2gAIAQEAAT8Q+du3bty9thogZOTWOtCrmiL9pQZtxMjdUjzHsYsWLFixYhE+b1f5mmUjSVZOPgxoZAUgNsWHhfea2ODl5NxyOvSEnxUqrQhpJB+Pd1+HGFDPKLjdd0unk16AE+KlVAgjQ3B+Pd1paymhWQsCSZIk3rlYBYoLGqIPhRkERJGphAmwGF4AePQ1PipVQI7abh/Hu6+hiNxVyEggCclloC0Uy28xkhQ5GBoRwksJHgw+FGibKXeR/VG/0VKqNnbXcP493WlgqEWNDdA0WNCrFiUWgoxNUDMSJi0zGwdlYggaDBbZRRAgDgaANgA+FGz7jwgF+lTxRs40dw/j3de3p+h1WNpoebi0r8TDGYZDU1Ye7HbOlIZSTiUvHOXXt6fo2VDdweLi8p8XDGBg3E6ehEi9D2AO4HgZNFl+PhjEEjZHHAXj2JQjIoyqyr8fDGT5NpVgmEQZpIzlCVXKv/LsMbDhjYcMbDhjYcMbDhjYcMaX8FMpC3R7EuD28KKRISyPQJrAPC/ftZQ1IiqcgU8tMUKzLBr7tGH2sMOqW0SEjhITh6NioHhn19n2ZvzrIiRdlc5zsLRi+xdJO6GdgsxMM8T4gEYbgBohlSQ+zhasBNZRO/fLs6JuEOUtitvRDdKcIJETInWNSDKeLd1ZbKINUmClYBRJLRQBoAShUZAQAAO9YANANPS1hP3NwCsWyDL2cNh5xcFm5S8xrFBbySEoTk1EsiJboJkfZme4IZsV8Ok7SFaHpVYdSbXWSrMiaC2XwcGgEAYKxn8qPGvcR3Ld9bNfWmCvYmzGQa1nCeolBzczr7OFxp9VMPoMaOG0NEe9Wv8A7HIlkuW9SZ3Tl4RP84o4/j7Rq7qxDpQ+0GTyE4eSlkRTAp2DXwUGzk8KTRyKTBRkj6WlayAuo0RYxaIjcsOaxY5FF4hIYHYYLwIE2Lq9GM7Rg5H8J+3tYQy1rHizJN4HhQTRKAuynlkdlNyFiXuwCehYkR28gVpYN7ug8V/x5p42B7TJP2o7KA0IsMQAFgNurB+5cqJw3s6N6R1/Vigi4A0CaU/dBAJKwkxJMbntZhfke4iU+I0qJ5QqlOXgj6soqrNupZJQWmMEEPgf/9k="></img>'
PAGE_STYLES = '''
:root {
    font-size: 20px;
}

*, ::file-selector-button {
    color: #293984;
    font-family: sans-serif;
    font-size: 1rem;
}

main {
    width: fit-content;
    max-width: 900px;
    margin: auto;
    margin-top: 3rem;
}

main img#logo {
    display: block;
    width: 75px;
    margin: auto;
    margin-bottom: 3rem;
}

button, input[type="submit"], ::file-selector-button {
    transition: background-color 0.2s, color 0.2s;
    border: 2px solid #293984;
    border-radius: 0.5rem;
    background-color: transparent;
    padding: 1rem;
    cursor: pointer;
}

::file-selector-button {
    margin-right: 1rem;
}

button:hover, input[type="submit"]:hover, ::file-selector-button:hover {
    background-color: #293984;
    color: white;
}

form p:last-of-type {
    margin-bottom: 1.5rem;
}

form * {
    display: block;
    width: fit-content;
    margin: auto;
    margin-top: 1rem;
}
'''
SPACE_RE = re.compile(r' {2,}')  # two or more spaces

def normalize_runs_in_paragraph(p):
    """
    Collapse multiple spaces to one within each run,
    and ensure only a single space across run boundaries
    (without touching paragraph breaks or run styling).
    """
    prev_ended_with_space = False

    for run in p.runs:
        t = run.text

        # 1) collapse internal multiples within the run
        t = SPACE_RE.sub(' ', t)

        # 2) if previous run ended with a space and this run starts with spaces,
        # squeeze to exactly one leading space on this run
        if prev_ended_with_space:
            # if there is any leading whitespace, reduce to exactly one regular space
            if t.startswith(' '):
                t = t.lstrip()
                t = ' ' + t
            # if it doesn't start with a space, we leave it as-is (no gap intended)

        run.text = t

        # 3) update boundary flag for next run
        prev_ended_with_space = run.text.endswith(' ')

def normalize_document_spaces(doc):
    # Paragraphs
    for p in doc.paragraphs:
        normalize_runs_in_paragraph(p)

    # Tables (iterate cells -> paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    normalize_runs_in_paragraph(p)

    # (Optional) headers/footers, if you want them too:
    """for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header is not None:
                for p in header.paragraphs:
                    normalize_runs_in_paragraph(p)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer is not None:
                for p in footer.paragraphs:
                    normalize_runs_in_paragraph(p)"""


def lambda_handler(event, context):
    print(event)

    resource = event.get("resource") or event.get("rawPath") or ""
    method = event.get("httpMethod") or event.get("requestContext", {}).get("http", {}).get("method", "")

    if resource == '/wordDocGenerator/chunk':  # Deprecated (not used in Upload/Fetch/Edit/Download Sequence)
        return handle_doc_chunk_upload(event)
    elif resource == '/wordDocGenerator/fileUpload':  # Upload file
        if method == 'GET':  # Serve web page
            return handle_file_upload_page(event)
        elif method == 'POST':  # Handle file upload
            return handle_file_upload(event)
    elif resource == '/wordDocGenerator/fileUploadData':  # Get JSON of uploaded .docx file
        return handle_file_upload_data(event)
    elif resource == '/wordDocGenerator/editFileUpload':  # Apply edits to uploaded .docx file
        return handle_edit_file(event)
    elif resource == '/wordDocGenerator/document':  # Download final document
        return handle_doc_download_page(event)


def new_s3_key():
    return str(uuid.uuid4())


def handle_doc_chunk_upload(event):
    body = json.loads(event['body'])

    s3 = boto3.client('s3')

    if 'upload_id' in body:  # If this is continuing an upload
        s3_key = body['upload_id']
        s3 = boto3.client('s3')
        s3_object = s3.get_object(Bucket=TMP_DATA_UPLOAD_BUCKET_NAME, Key=s3_key)
        doc_text = s3_object['Body'].read().decode('utf-8')

        curr_doc_data = json.loads(doc_text)
    else:  # If this is the first chunk
        s3_key = new_s3_key()

        curr_doc_data = {
            'doc_data': {
                'paragraphs': []
            }
        }

    # Update the filename unless it's not given (it might have been specified on an earlier chunk)
    if 'filename' in body:
        curr_doc_data['filename'] = body['filename']

    curr_doc_data['doc_data']['paragraphs'].extend(body['doc_data']['paragraphs'])

    s3.put_object(
        Bucket=TMP_DATA_UPLOAD_BUCKET_NAME,
        Key=s3_key,
        Body=json.dumps(curr_doc_data).encode('utf-8'),
        ContentType='application/json'
    )

    return {
        "statusCode": 200,
        "headers": {
            "Content-Type": "application/json",
            "Access-Control-Allow-Origin": "*"
        },
        "body": json.dumps({
            "upload_id": s3_key
        })
    }


def handle_file_upload_page(event):
    add_query_params_def = '''
    document.addEventListener('DOMContentLoaded', function() {
        const form = document.querySelector('form');
        const urlParams = new URLSearchParams(window.location.search);
        const upload_id = urlParams.get('upload_id'); // Get the upload ID from the query params

        form.action = `${window.location.pathname}?upload_id=${upload_id}`;
    });
    '''
    return {
        "statusCode": 200,
        "headers": {
            "Content-Type": "text/html",
            "Access-Control-Allow-Origin": "*"
        },
        "body": f"""<html>
            <head>
                <title>Upload Document</title>
                <style>
                    {PAGE_STYLES}
                </style>
                {FAVICON_ELT}
            </head>
            <body>
                <main>
                    {LOGO_ELT}
                    <form action="{event['requestContext']['path']}" method="post" enctype="multipart/form-data">
                        <p>Upload your file to the GPT using the input below.</p>
                        <p style="font-size: 0.6rem; margin-top: 0.1rem">This link is for one chat only. If you did not reach this page through the GPT, please ask the GPT for a new link.</p>
                        <input type="file" name="file" accept=".docx, .pdf" />
                        <input type="submit" value="Upload" />
                    </form>
                </main>
                <script>{add_query_params_def}</script>
            </body>
        </html>
        """
    }


def handle_file_upload(event):
    print("Handling file upload...")

    # Safely extract query params and upload_id
    query_params = event.get('queryStringParameters') or {}
    upload_id = query_params.get('upload_id')

    if not upload_id:
        print("❌ Missing upload_id in query params")
        return {
            "statusCode": 400,
            "body": "Missing upload_id query parameter"
        }

    # Decode body
    body = event.get("body")
    if body is None:
        return {
            "statusCode": 400,
            "body": "Missing request body"
        }

    if event.get("isBase64Encoded", False):
        try:
            body = base64.b64decode(body)
        except Exception as e:
            print(f"Base64 decode failed: {e}")
            return {"statusCode": 400, "body": "Invalid base64 encoding"}
    elif isinstance(body, str):
        body = body.encode("iso-8859-1")

    # Normalize header keys
    headers = {k.lower(): v for k, v in (event.get("headers") or {}).items()}
    content_type = headers.get("content-type")

    if not content_type:
        return {"statusCode": 400, "body": "Missing Content-Type header"}

    # Parse the multipart form data
    try:
        fs = cgi.FieldStorage(
            fp=io.BytesIO(body),
            headers={"content-type": content_type},
            environ={"REQUEST_METHOD": "POST"}
        )
    except Exception as e:
        print(f"Form parsing failed: {e}")
        return {"statusCode": 400, "body": f"Error parsing form data: {str(e)}"}

    if "file" not in fs:
        print("❌ No file found in upload")
        return {"statusCode": 400, "body": "No file uploaded"}

    uploaded_file = fs["file"]
    file_content = uploaded_file.file.read()
    FILE_EXTENSION = ""
    if uploaded_file.filename:
        _, ext = os.path.splitext(uploaded_file.filename)
        FILE_EXTENSION = ext.lower()  # e.g. ".pdf" or ".docx"
    output_stream = None
    if FILE_EXTENSION == ".pdf":
        CONTENT_TYPE = "application/pdf"
    elif FILE_EXTENSION == ".docx":
        CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        doc = Document(io.BytesIO(file_content))
        normalize_document_spaces(doc)

        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
    # Upload to S3
    s3 = boto3.client("s3")
    s3_key = upload_id
    s3.put_object(
        Bucket=TMP_DOC_UPLOAD_BUCKET_NAME,
        Key=s3_key,
        Body=(output_stream.getvalue() if output_stream else file_content),
        ContentType=CONTENT_TYPE,
    )

    print(f"✅ Uploaded file to S3 with key: {s3_key}")

    # Return confirmation page
    return {
        "statusCode": 200,
        "headers": {
            "Content-Type": "text/html",
            "Access-Control-Allow-Origin": "*"
        },
        "body": f'''<html>
        <head>
            <title>Upload Document</title>
            <style>{PAGE_STYLES}</style>
            {FAVICON_ELT}
        </head>
        <body>
            <main>
                {LOGO_ELT}
                <p>Document uploaded successfully. Please return to the GPT and type "done" in the chat.</p>
            </main>
        </body>
        </html>
        '''
    }


def handle_file_upload_data(event):
    query_params = event.get('queryStringParameters') or {}

    s3_key = query_params['upload_id']
    cursor = query_params['cursor']
    size = query_params['size']

    s3 = boto3.client('s3')
    s3_object = s3.get_object(Bucket=TMP_DOC_UPLOAD_BUCKET_NAME, Key=s3_key)
    file_content = s3_object['Body'].read()
    file_stream = io.BytesIO(file_content)

    doc = Document(file_stream)

    # Call helper function from word_doc.py
    result = json_data_from_doc(doc, cursor, size)

    # Transform to expected API format
    paragraphs = result['doc_data']['paragraphs']
    next_cursor = result['next-cursor']

    # Flatten paragraphs into chunks
    chunks = []
    for idx, para in enumerate(paragraphs):
        chunks.append({
            "kind": "paragraph",
            "runs": para[0]['runs'] if para else [],
            "chunk_idx": int(cursor) + idx
        })

    return {
        "statusCode": 200,
        "headers": {
            "Content-Type": "application/json",
            "Access-Control-Allow-Origin": "*"
        },
        "body": json.dumps({
            "data": chunks,
            "cursor": int(next_cursor) if next_cursor != "STOP" else int(cursor) + len(chunks),
            "more_data": next_cursor != "STOP",
            "total_chunks": len(doc.paragraphs)
        })
    }


# This function is unable to read existing insertions and deletions, so they should be skipped
def handle_edit_file(event):
    body = json.loads(event['body'])

    s3_key = body['upload_id']

    s3 = boto3.client('s3')
    s3_object = s3.get_object(Bucket=TMP_DOC_UPLOAD_BUCKET_NAME, Key=s3_key)
    file_content = s3_object['Body'].read()
    file_stream = io.BytesIO(file_content)

    doc = Document(file_stream)

    author = body['author']
    print(body)
    change_id_counter = 1
    for e in body['edits']:
        edit = e["edit_spec"]
        print(e)
        edit_type = edit['type']
        comment = edit.get('comment', None)
        if edit_type == 'replace':
            surrounding_text = edit['surrounding_text']
            find = edit['find']
            replace = edit['replace']

            if find not in surrounding_text:
                return {
                    "statusCode": 400,
                    "headers": {
                        "Content-Type": "application/json",
                        "Access-Control-Allow-Origin": "*"
                    },
                    "body": json.dumps({
                        'problem': 'Value of "surrounding_text" does not contain value of "find".',
                        'surrounding_text': surrounding_text,
                        'find': find,
                        'consequence': 'No edits were applied.',
                        'tip': 'Ensure value of "surrounding_text" contains the value of "find"'
                    })
                }

            # Determine indices of edits
            curr_edit = None
            text_so_far = ''
            found_edit_location = False
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    text_so_far += run.text

                if surrounding_text in text_so_far:  # This and previous paragraphs contain the edit
                    start_ind = text_so_far.index(surrounding_text) + surrounding_text.index(find)
                    end_ind = text_so_far.index(surrounding_text) + surrounding_text.index(find) + len(find)

                    curr_edit = (start_ind, end_ind, replace)
                    found_edit_location = True
                    break  # After this, text_so_far will always contain surrounding_text
            
            print(surrounding_text in text_so_far)

            if not found_edit_location:
                return {
                    "statusCode": 400,
                    "headers": {
                        "Content-Type": "application/json",
                        "Access-Control-Allow-Origin": "*"
                    },
                    "body": json.dumps({
                        'problem': 'No section of the document matched surrounding_text exactly.',
                        'surrounding_text': surrounding_text,
                        'consequence': 'No edits were applied.',
                        'tip': 'Make sure that the surrounding text matches the document EXACTLY (spaces, punctuation, spelling, grammar, etc.) and try again WITHOUT SPEAKING'
                    })
                }

            # Apply the edit now that the indices have been found
            chars_so_far = 0
            found_start = False
            edit_runs = []
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    next_chars_so_far = chars_so_far + len(run.text)  # Chars so far after adding this run
                    if next_chars_so_far > curr_edit[0] and not found_start:  # If this is the first run being targeted
                        _, run = split_run(run, curr_edit[0] - chars_so_far)
                        chars_so_far = curr_edit[0]
                        found_start = True
                    elif curr_edit[0] < chars_so_far < next_chars_so_far < curr_edit[
                        1]:  # If this run is in the middle of what's being targeted
                        pass  # The logic below will still incorporate the run

                    if next_chars_so_far >= curr_edit[1]:  # If this is the last run being targeted
                        run, _ = split_run(run, curr_edit[1] - chars_so_far)

                        change_id_counter += 1

                        edit_runs.append(run)
                        break

                    if found_start:
                        edit_runs.append(run)

                    chars_so_far = next_chars_so_far

                if found_start:
                    break

            del_elem = wrap_runs_with_tracked_deletion(edit_runs, author, change_id_counter)
            change_id_counter += 1
            ins_elem = create_tracked_insertion(del_elem, replace, author, change_id_counter)
            change_id_counter += 1

            if comment:
                add_comment_to = ins_elem if len(replace) > 0 else del_elem
                comment_run_elts = add_comment_to.xpath('.//w:r', namespaces=add_comment_to.nsmap)
                doc.add_comment(
                    runs=[Run(r_elt, paragraph) for r_elt in comment_run_elts],
                    text=comment,
                    author=author
                )

        elif edit_type == 'insert_text':
            adjacent_text = edit.get('adjacent_text') or edit.get('surrounding_text')
            insert_pos = edit['insert_pos']
            insert_runs = edit['insert']

            # Store ordered map of runs and corresponding text
            run_texts = []  # Format: (run, text)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run_texts.append((run, run.text))

            found_edit_location = False
            edit_run = None
            edit_run_split_ind = None
            if insert_pos == 'before':
                # Find the first run that is part of adjacent_text and store as edit_run
                # Also find the index at which to split edit_run, and store as edit_run_split_ind
                future_text = ''.join([text for (run, text) in run_texts])
                while adjacent_text in future_text:
                    found_edit_location = True

                    # These will get overridden if this is not the first run part of adjacent_text
                    edit_run_split_ind = future_text.index(adjacent_text)  # Future text starts with edit run
                    edit_run = run_texts.pop(0)[0]  # Remove and store first run

                    # Update future text
                    future_text = ''.join([text for (run, text) in run_texts])

            elif insert_pos == 'after':
                # Find the last run that is part of adjacent_text and store as edit_run
                # Also find the index at which to split edit_run, and store as edit_run_split_ind
                prev_text_until_here = ''
                for i in range(len(run_texts)):
                    text_until_here = ''.join([text for (run, text) in run_texts[:i]])
                    if adjacent_text in text_until_here:
                        found_edit_location = True

                        run, text = run_texts[i - 1]  # Get the data for the last run
                        edit_run = run
                        edit_run_split_ind = text_until_here.index(adjacent_text) + len(adjacent_text) - len(
                            prev_text_until_here)
                        break

                    prev_text_until_here = text_until_here

            if not found_edit_location:
                return {
                    "statusCode": 400,
                    "headers": {
                        "Content-Type": "application/json",
                        "Access-Control-Allow-Origin": "*"
                    },
                    "body": json.dumps({
                        'problem': 'No section of the document matched adjacent_text exactly.',
                        'adjacent_text': adjacent_text,
                        'consequence': 'No edits were applied.',
                        'tip': 'Make sure that the adjacent text matches the document EXACTLY (spaces, punctuation, spelling, grammar, etc.) and try again WITHOUT SPEAKING'
                    })
                }

            prev_run, next_run = split_run(edit_run, edit_run_split_ind)
            ins_elem = create_tracked_insertion(prev_run._element, insert_runs, author, change_id_counter)
            change_id_counter += 1

            if comment:
                add_comment_to = ins_elem
                comment_run_elts = add_comment_to.xpath('.//w:r', namespaces=add_comment_to.nsmap)
                doc.add_comment(
                    runs=[Run(r_elt, paragraph) for r_elt in comment_run_elts],
                    text=comment,
                    author=author
                )

        """elif edit_type == 'insert_paragraph':
            adjacent_text = edit.get('adjacent_text') or edit.get('surrounding_text')
            insert_pos = edit['insert_pos']
            insert_runs = edit['insert']
            is_list_item = edit.get('is_list_item', False)
            list_level = edit.get('list_level', 0)  # NEW: support sub-bullets

            # Store ordered map of paragraphs and corresponding text
            paragraph_texts = []  # Item format: (paragraph, text)
            for paragraph in doc.paragraphs:
                curr_p_text = ''  # Keep track of this paragraph's text
                for run in paragraph.runs:
                    curr_p_text += run.text
                paragraph_texts.append((paragraph, curr_p_text))

            insert_ind = None
            found_insert_location = False

            if insert_pos == 'before':
                for i, (p, text) in enumerate(paragraph_texts):
                    future_text = ''.join([item[1] for item in paragraph_texts[i:]])
                    if future_text.startswith(adjacent_text):
                        insert_ind = i
                        found_insert_location = True
                        break
            elif insert_pos == 'after':
                for i, (p, text) in enumerate(paragraph_texts):
                    if text.endswith(adjacent_text):
                        insert_ind = i + 1
                        found_insert_location = True
                        break

            if not found_insert_location:
                return {
                    "statusCode": 400,
                    "headers": {
                        "Content-Type": "application/json",
                        "Access-Control-Allow-Origin": "*"
                    },
                    "body": json.dumps({
                        'problem': f'Could not find a paragraph {"starting" if insert_pos == "before" else "ending"} with adjacent_text',
                        'adjacent_text': adjacent_text,
                        'insert_pos': insert_pos,
                        'consequence': 'No edits were applied.',
                        'tip': 'Make sure that you are'
                    })
                }

            new_paragraph = doc.paragraphs[insert_ind].insert_paragraph_before()

            # ---------- NEW LIST HANDLING ----------
            if is_list_item:
                from copy import deepcopy
                from lxml import etree
                nsmap = new_paragraph._element.nsmap

                def apply_list_formatting(new_paragraph, list_level):
                    # Try to find a matching list paragraph at the same level
                    for p in doc.paragraphs:
                        ppr = p._element.pPr
                        if ppr is not None:
                            numPr = ppr.find('.//w:numPr', nsmap)
                            if numPr is not None:
                                ilvl = numPr.find('.//w:ilvl', nsmap)
                                if ilvl is not None and int(ilvl.get('{%s}val' % nsmap['w'])) == list_level:
                                    new_paragraph._element.insert(0, deepcopy(ppr))
                                    return

                    # Otherwise, build a fresh numPr with desired level
                    numPr = etree.Element('{%s}numPr' % nsmap['w'])
                    ilvl = etree.Element('{%s}ilvl' % nsmap['w'])
                    ilvl.set('{%s}val' % nsmap['w'], str(list_level))
                    numId = etree.Element('{%s}numId' % nsmap['w'])
                    numId.set('{%s}val' % nsmap['w'], "1")  # assumes list definition ID=1

                    numPr.append(ilvl)
                    numPr.append(numId)

                    if new_paragraph._element.pPr is None:
                        new_paragraph._element.insert(0, etree.Element('{%s}pPr' % nsmap['w']))
                    new_paragraph._element.pPr.append(numPr)

                apply_list_formatting(new_paragraph, list_level)
            # ---------- END LIST HANDLING ----------

            dummy_run = new_paragraph.add_run()  # Reference point for tracked insertion
            ins_elem = create_tracked_insertion(dummy_run._element, insert_runs, author, change_id_counter)
            change_id_counter += 1

            if comment:
                add_comment_to = ins_elem
                comment_run_elts = add_comment_to.xpath('.//w:r', namespaces=add_comment_to.nsmap)
                doc.add_comment(
                    runs=[Run(r_elt, new_paragraph) for r_elt in comment_run_elts],
                    text=comment,
                    author=author
                )
        """
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    s3.upload_fileobj(
        buffer,
        TMP_DOC_UPLOAD_BUCKET_NAME,
        s3_key,
        ExtraArgs={'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'}
    )

    return {
        "statusCode": 200,
        "headers": {
            "Content-Type": "application/json",
            "Access-Control-Allow-Origin": "*"
        },
        "body": json.dumps({'upload_id': s3_key})
    }


def handle_doc_download_page(event):
    query_params = event.get('queryStringParameters', {})
    s3_key = query_params['upload_id']
    doc_source = query_params['source']

    s3 = boto3.client('s3')

    if doc_source == 'generated':
        s3_object = s3.get_object(Bucket=TMP_DATA_UPLOAD_BUCKET_NAME, Key=s3_key)
        doc_bytes = s3_object['Body'].read()
        doc_json = json.loads(doc_bytes.decode('utf-8'))
        filename = doc_json.get('filename', 'document.docx')
        encoded_doc = doc_from_json_doc_data(doc_json['doc_data'])
    elif doc_source == 'uploaded':
        s3_object = s3.get_object(Bucket=TMP_DOC_UPLOAD_BUCKET_NAME, Key=s3_key)
        doc_bytes = s3_object['Body'].read()
        filename = query_params['filename']
        encoded_doc = base64.b64encode(doc_bytes).decode("utf-8")

    doc_uri = f'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{encoded_doc}'

    return {
        "statusCode": 200,
        "headers": {
            "Content-Type": "text/html",
            "Access-Control-Allow-Origin": "*"
        },
        "body": f'''<html>
            <head>
                <title>Download Word File</title>
                {FAVICON_ELT}
                <style>
                    {PAGE_STYLES}
                </style>
            </head>
            <body>
                <main>
                    {LOGO_ELT}
                    <p><a href="{doc_uri}" download="{filename}">Click here</a> to download your Word file.</p>
                </main>
            </body>
            </html>
        '''
    }
